import os
import base64
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
import requests
import json
import time

from docx import Document
from docx.shared import Inches

# 设置千问API的密钥和基础URL
os.environ["QWEN_API_KEY"] = "sk-56dd39ce661c4174bdc53f93b7221509"
os.environ["QWEN_BASE_URL"] = "https://dashscope.aliyuncs.com/api/v1"


def generate_ecg_details(ecg_image):
    """使用千问VL-max模型分析ECG图像并生成中文报告"""
    # 打开并处理图像
    try:
        image = Image.open(ecg_image)
        if image.mode != 'RGB':
            if image.mode == 'RGBA':
                bg = Image.new("RGB", image.size, (255, 255, 255))
                bg.paste(image, (0, 0), image)
                image = bg
            else:
                image = image.convert('RGB')

        buffered = BytesIO()
        image.save(buffered, format="JPEG")
        img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
    except Exception as e:
        st.error(f"处理图像失败: {str(e)}")
        return "## 报告生成失败\n错误: 无法处理上传的心电图图像"

    current_date = datetime.now().strftime('%Y-%m-%d')

    prompt = """## 角色：您是一名专业的心脏病专家
    ## 任务：
    请详细分析这张心电图（ECG）图像并提供全面的报告。请根据从图像中提取的信息填写所有字段。
    如果完全无法确定某项信息，请注明"无法从提供的ECG图像确定"。请不要使用占位符如'[待填写]'。
    在可能的情况下做出合理推断，但要明确指出这些是推测内容。请严格按照以下结构编写报告：
    ---
    ### 心电图分析报告

    ### 1. 患者基本信息
    ◦ 姓名:

    ◦ 年龄:

    ◦ 性别:

    ◦ 身份证号:

    ◦ ECG检查日期:


    ### 2. 临床信息
    ◦ ECG检查原因:

    ◦ 相关病史:

    ◦ 当前用药:


    ### 3. ECG技术参数
    ◦ 使用的心电图机型号:

    ◦ 导联配置:

    ◦ 校准情况:

    ◦ 记录质量:


    ### 4. ECG主要发现
    节律与心率
    ◦ 心率:

    ◦ 节律:

    ◦ P波特征:

    ◦ PR间期:

    ◦ QRS波群:

    ◦ QT/QTc间期:

    ◦ ST段:

    ◦ T波特征:


    心电轴
    ◦ P波电轴:

    ◦ QRS电轴:

    ◦ T波电轴:


    传导与形态学
    ◦ 心房传导:

    ◦ 心室传导:

    ◦ QRS形态:

    ◦ ST-T改变:


    ### 5. 分析结论
    ◦ 正常/异常:

    ◦ 主要诊断/发现:

    ◦ 是否存在心率失常:


    ### 6. 总结与建议
    ◦ 结论概要:

    ◦ 医疗建议:



    ### 7. 报告医师
    ◦ 医师姓名:

    ◦ 签名: 

    ◦ 报告日期: {}

    ---
    ## 要求：
    1. 只输出报告内容，不要添加任何额外说明
    2. 使用Markdown格式严格按照上述结构输出
    3. 保持专业医疗报告风格""".format(current_date)

    headers = {
        "Authorization": f"Bearer {os.environ['QWEN_API_KEY']}",
        "Content-Type": "application/json",
    }

    # 同步API端点
    api_endpoint = f"{os.environ['QWEN_BASE_URL']}/services/aigc/multimodal-generation/generation"

    content = [
        {"image": f"data:image/jpeg;base64,{img_str}"},
        {"text": prompt}
    ]

    payload = {
        "model": "qwen-vl-max",
        "input": {
            "messages": [
                {
                    "role": "user",
                    "content": content
                }
            ]
        },
        "task": "multimodal-generation",  # 添加必需的task参数
        "parameters": {
            "temperature": 0.1,  # 更保守的温度设置以提高准确性
            "top_p": 0.8,
            "max_tokens": 3000
        }
    }

    try:
        start_time = time.time()
        response = requests.post(api_endpoint, headers=headers, json=payload, timeout=60)
        response.raise_for_status()  # 检查HTTP错误

        result = response.json()

        # 调试输出
        st.write("API响应:", json.dumps(result, indent=2, ensure_ascii=False))

        # 检查API响应结构
        if "output" not in result or "choices" not in result["output"]:
            error_info = result.get("error", result)
            st.error(f"API返回格式不正确: {json.dumps(error_info, indent=2)}")
            return "## 报告生成失败\n错误：API返回格式不正确"

        # 提取生成的报告内容 - 处理可能的复杂结构
        report_content = result["output"]["choices"][0]["message"]["content"]

        # 处理不同的返回格式
        if isinstance(report_content, str):
            # 如果是字符串，直接返回
            return report_content
        elif isinstance(report_content, list):
            # 如果是列表，提取所有文本内容
            text_parts = []
            for item in report_content:
                if isinstance(item, dict) and "text" in item:
                    text_parts.append(item["text"])
                elif isinstance(item, str):
                    text_parts.append(item)
            return "\n".join(text_parts)
        elif isinstance(report_content, dict) and "text" in report_content:
            # 如果是字典且包含text字段
            return report_content["text"]
        else:
            # 其他未知格式，转换为字符串
            return str(report_content)

    except requests.exceptions.HTTPError as err:
        error_detail = ""
        if hasattr(response, 'text'):
            error_detail = f"状态码: {response.status_code}\n响应内容: {response.text[:500]}"
        st.error(f"千问API调用失败: {err}\n{error_detail}")
        return f"## 报告生成失败\n错误信息: HTTP {response.status_code if 'response' in locals() else 'N/A'} - {err}"
    except (KeyError, IndexError, TypeError) as e:
        error_detail = f"解析API响应失败: {e}"
        if 'response' in locals() and hasattr(response, 'text'):
            error_detail += f"\n响应内容: {response.text[:500]}"
        st.error(error_detail)
        return "## 报告生成失败\n错误：无法解析AI返回内容。"
    except Exception as e:
        st.error(f"处理过程中发生意外错误: {str(e)}")
        return f"## 报告生成失败\n错误信息: {str(e)}"


def create_doc(report_text, ecg_image):
    """将报告转换为Word文档"""
    doc = Document()
    doc.add_heading('心电图分析报告', 0)  # 文档主标题

    # 添加报告基本信息
    p = doc.add_paragraph()
    p.add_run("报告生成时间: ").bold = True
    p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    # 处理报告内容
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # 处理标题
        if line.startswith('### '):
            doc.add_heading(line[4:], level=1)
        elif line.startswith('◦ '):
            # 处理列表项
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # 普通段落
            doc.add_paragraph(line)

    # 添加原始心电图
    doc.add_heading('心电图影像', level=1)
    ecg_image.seek(0)  # 重置文件指针
    image_data = ecg_image.read()  # 读取图像数据
    image_stream = BytesIO(image_data)  # 创建新的BytesIO对象
    doc.add_picture(image_stream, width=Inches(6))

    # 添加免责声明
    doc.add_heading('免责声明', level=1)
    doc.add_paragraph(
        "本报告由AI系统(MED360)生成，仅供参考使用。准确诊断和治疗请咨询专业医疗人员。由于图像质量和信息限制，本报告不构成医疗诊断依据。如发现任何异常，请立即进行专业医疗评估。")
    doc.add_paragraph("生成时间: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    # 将文档保存到内存中
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def main():
    """应用主界面"""
    st.set_page_config(
        page_title="心电图分析系统",
        page_icon="❤️",
        layout="centered",
        initial_sidebar_state="collapsed"
    )

    # 标题区域
    st.title("❤️USST智能医学实验室ECG——AI分析系统")
    st.markdown("心电图智能分析工具")
    st.divider()

    # 创建三个步骤的UI
    tab1, tab2, tab3 = st.tabs(["上传心电图", "生成报告", "下载报告"])

    with tab1:
        st.subheader("第一步: 上传心电图")
        ecg_image = st.file_uploader(
            "请选择心电图图像文件 (PNG/JPG/JPEG格式)",
            type=["png", "jpg", "jpeg"],
            key="ecg_uploader",
            help="上传清晰的心电图扫描或照片"
        )

        if ecg_image:
            st.success("心电图图像已上传成功!")
            st.image(ecg_image, caption='您上传的心电图', width=300)

            # 存储到会话状态
            st.session_state.ecg_image = ecg_image
            st.session_state.ready_for_step2 = True

    with tab2:
        st.subheader("第二步: 生成分析报告")

        if not st.session_state.get("ready_for_step2", False):
            st.warning("请先上传心电图图像")
            return

        if st.button("✨ 开始AI分析", type="primary", use_container_width=True):
            with st.status("AI正在分析心电图...", expanded=True) as status:
                st.write("解析心电图图像数据...")
                time.sleep(1)

                st.write("调用USST智能医学实验室模型进行分析...")
                ecg_image = st.session_state.ecg_image
                ecg_details = generate_ecg_details(ecg_image)

                st.write("正在生成报告...")
                time.sleep(0.5)

                # 存储结果到会话状态
                st.session_state.ecg_details = ecg_details
                st.session_state.report_generated = True
                status.update(label="分析完成!", state="complete")

            # 显示报告预览
            st.subheader("心电图分析预览")
            with st.expander("查看完整报告", expanded=True):
                st.markdown(ecg_details)

    with tab3:
        st.subheader("第三步: 下载报告")

        if not st.session_state.get("report_generated", False):
            st.info("请先生成分析报告")
            return

        # 创建并下载Word报告
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("完整报告(Word)")
            st.info("包含完整分析和心电图图像的专业报告")

            try:
                with st.spinner("正在准备Word文档..."):
                    doc_file_stream = create_doc(
                        st.session_state.ecg_details,
                        st.session_state.ecg_image
                    )

                st.download_button(
                    label="📥 下载Word报告",
                    data=doc_file_stream,
                    file_name=f"ECG报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"文档创建失败: {str(e)}")
                st.error("如果多次失败，请尝试下载文本报告")

        with col2:
            st.subheader("文本报告(TXT)")
            st.info("仅包含分析文本的简约格式")

            # 确保ecg_details是字符串
            report_text = st.session_state.ecg_details
            if not isinstance(report_text, str):
                if isinstance(report_text, list):
                    report_text = "\n".join(report_text)
                else:
                    report_text = str(report_text)

            text_report = BytesIO(report_text.encode('utf-8'))
            st.download_button(
                label="📄 下载文本报告",
                data=text_report,
                file_name=f"ECG报告文本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )

        # 添加反馈区域
        st.divider()
        st.subheader("您的反馈")
        feedback = st.text_area("您的意见对我们改进系统非常有价值")
        if st.button("提交反馈", use_container_width=True):
            st.success("感谢您的反馈!我们会不断优化分析质量")




# 初始化会话状态
if 'ready_for_step2' not in st.session_state:
    st.session_state.ready_for_step2 = False

if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False

if __name__ == '__main__':
    main()  # 启动应用